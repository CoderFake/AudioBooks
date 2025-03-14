import os
import re
import tempfile
import logging
import asyncio
import subprocess
from typing import Optional, Dict, Any, List
from pathlib import Path

import docx
import PyPDF2
import textract
import chardet
from bs4 import BeautifulSoup

from utils.text_processor import preprocess_text

logger = logging.getLogger(__name__)


async def read_file_content(file_path: str) -> Optional[str]:
    try:
        file_extension = Path(file_path).suffix.lower()

        if file_extension == '.txt':
            return await read_text_file(file_path)
        elif file_extension == '.docx':
            return await read_docx_file(file_path)
        elif file_extension == '.pdf':
            return await read_pdf_file(file_path)
        elif file_extension in ['.html', '.htm']:
            return await read_html_file(file_path)
        elif file_extension == '.md':
            return await read_markdown_file(file_path)
        elif file_extension == '.epub':
            return await read_epub_file(file_path)
        elif file_extension in ['.doc', '.rtf', '.odt']:
            return await read_with_textract(file_path)
        else:
            try:
                return await read_text_file(file_path)
            except:
                logger.warning(f"Unsupported file format: {file_extension}")
                return None
    except Exception as e:
        logger.exception(f"Error reading file: {str(e)}")
        return None


async def read_text_file(file_path: str) -> str:

    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        return content
    except UnicodeDecodeError:
        with open(file_path, 'rb') as f:
            raw_data = f.read()

        detected = chardet.detect(raw_data)
        encoding = detected['encoding']

        with open(file_path, 'r', encoding=encoding) as f:
            content = f.read()
        return content


async def read_docx_file(file_path: str) -> str:
    doc = docx.Document(file_path)
    full_text = []

    for para in doc.paragraphs:
        full_text.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    full_text.append(para.text)

    return '\n'.join(full_text)


async def read_pdf_file(file_path: str) -> str:
    with open(file_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        full_text = []

        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            full_text.append(page.extract_text())

    return '\n'.join(full_text)


async def read_html_file(file_path: str) -> str:
    with open(file_path, 'r', encoding='utf-8') as f:
        html_content = f.read()

    soup = BeautifulSoup(html_content, 'html.parser')

    for script in soup(["script", "style"]):
        script.extract()

    text = soup.get_text()

    lines = (line.strip() for line in text.splitlines())
    chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
    text = '\n'.join(chunk for chunk in chunks if chunk)

    return text


async def read_markdown_file(file_path: str) -> str:
    return await read_text_file(file_path)


async def read_epub_file(file_path: str) -> str:
    content = textract.process(file_path).decode('utf-8')
    return content


async def read_with_textract(file_path: str) -> str:
    content = textract.process(file_path).decode('utf-8')
    return content


async def process_uploaded_file(file_path: str, preprocess: bool = True) -> Dict[str, Any]:
    try:
        file_info = {
            "file_path": file_path,
            "file_name": os.path.basename(file_path),
            "file_size": os.path.getsize(file_path),
            "file_extension": os.path.splitext(file_path)[1].lower(),
        }

        content = await read_file_content(file_path)
        if content is None:
            raise ValueError(f"Cannot read content from file: {file_path}")

        if preprocess:
            content = preprocess_text(content)

        file_info["content"] = content
        file_info["word_count"] = len(content.split())
        file_info["char_count"] = len(content)

        return file_info
    except Exception as e:
        logger.exception(f"Error processing uploaded file: {str(e)}")
        raise


async def extract_content_by_pages(file_path: str) -> List[Dict[str, Any]]:

    try:
        file_extension = Path(file_path).suffix.lower()

        if file_extension == '.pdf':
            return await extract_pdf_pages(file_path)
        elif file_extension in ['.docx']:
            return await extract_docx_pages(file_path)
        else:
            content = await read_file_content(file_path)
            if content is None:
                raise ValueError(f"Cannot read content from file: {file_path}")

            return [{
                "index": 0,
                "content": content,
                "word_count": len(content.split()),
                "char_count": len(content)
            }]
    except Exception as e:
        logger.exception(f"Error extracting content by pages: {str(e)}")
        raise


async def extract_pdf_pages(file_path: str) -> List[Dict[str, Any]]:
    pages = []

    with open(file_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)

        for i in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[i]
            content = page.extract_text()

            pages.append({
                "index": i,
                "content": content,
                "word_count": len(content.split()),
                "char_count": len(content)
            })

    return pages


async def extract_docx_pages(file_path: str) -> List[Dict[str, Any]]:
    doc = docx.Document(file_path)
    pages = []

    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            pages.append({
                "index": i,
                "content": para.text,
                "word_count": len(para.text.split()),
                "char_count": len(para.text)
            })

    return pages


async def merge_text_files(file_paths: List[str], output_file: str) -> bool:
    try:
        with open(output_file, 'w', encoding='utf-8') as out_file:
            for file_path in file_paths:
                content = await read_file_content(file_path)
                if content:
                    out_file.write(content)
                    out_file.write('\n\n')

        return True
    except Exception as e:
        logger.exception(f"Error merging text files: {str(e)}")
        return False


async def convert_to_plain_text(file_path: str, output_file: str) -> bool:
    try:
        content = await read_file_content(file_path)
        if content is None:
            raise ValueError(f"Cannot read content from file: {file_path}")

        with open(output_file, 'w', encoding='utf-8') as out_file:
            out_file.write(content)

        return True
    except Exception as e:
        logger.exception(f"Error converting to plain text: {str(e)}")
        return False


async def detect_text_language(text: str) -> str:
    from langdetect import detect

    try:
        return detect(text)
    except:
        return 'vi'


async def count_pages(file_path: str) -> int:
    try:
        file_extension = Path(file_path).suffix.lower()

        if file_extension == '.pdf':
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                return len(pdf_reader.pages)
        elif file_extension == '.docx':
            doc = docx.Document(file_path)
            return len(doc.paragraphs)
        else:
            content = await read_file_content(file_path)
            if content:
                return len(content.splitlines())
            return 0
    except Exception as e:
        logger.exception(f"Error counting pages: {str(e)}")
        return 0